import json
from gpti import gpt
import fitz
from docx import Document
import io
from io import BytesIO
from flask import Response

from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfgen import canvas

# Contexto default utilizado para responder los examenes
contexto = [
    {
        "role": "assistant",
        "content": "Hola, como puedo ayudarte"
    },
    {
        "role": "user",
        "content": """Responde las preguntas de opcion multiple, regresa una respuesta en un diccionario de python que contenga lo siguiente:
                        - titulo
                        - lista_preguntas: cada pregunta es un diccionario con 'texto' y otra lista con las 4 'opciones' y su 'respuesta_correcta'
                    Tu respuesta debe iniciar y terminar solo con el diccionario, es decir \{\}. Utiliza las llaves de manera exacta como se describieron previamente
                    """
    }
]

def get_chatGPT_response(question: str, checado=False, context : list = contexto, model='gpt-4', md=False):
    """
    Hace una conexion a chatGPT para responder una pregunta en un determinado contexto y regresando
    el resultado en forma de diccionario

    Parametros:
    - context (str): contexto para el chat
    - course (str): materia a la cual pertenece el examen
    - question (str): preguntas del examen que se busca resolver
    - model (str): modelo de chatGPT a usar

    Regresa:
    - diccionario: Representa la informacion detallade del examen incluyendo la respuesta correcta

    """
    if checado:
        additional = {
            "role": "user",
            "content": """El documento que recibiras ya estará resuelto, por lo tanto, busca la respuesta y
                        ponla en su correspondiente sección. NO intentes resolverla.
                    """
        }
        context.append(additional)

    res = gpt(messages=context, prompt=question, model=model, markdown=md)

    if res.error != None:
        return json.dumps(res.error)
    else:
        return json.loads(json.dumps(res.result['gpt']))

def extract_text_from_file(content, filename):
    """
    Determina si un documento es de tipo pdf o docx y extrae su contenido

    Parametros:
    - content: el documento a analizar
    - filename: nombre del documento

    Regresa:
    - str: el texto perteneciente al documento
    """

    if filename.endswith('.pdf'):
        return extract_text_from_pdf(content)
    elif filename.endswith('.docx'):
        return extract_text_from_docx(content)
    else:
        return "Unsupported file format"

def extract_text_from_pdf(content):
    """
    Concatena el contenido de un documento .pdf en un str

    Parametros:
    - content: contenido del documento

    Regresa:
    - str: el contenido en una variable
    """

    text = ""
    with fitz.open(stream=content, filetype="pdf") as doc:
        for page in doc:
            text += page.get_text()
    return text

def extract_text_from_docx(content):
    """
    Convierte el contenido de un documento .docx en un str

    Parametros:
    - content: contenido del documento

    Regresa:
    - str: el contenido en una variable
    """

    doc = Document(io.BytesIO(content))
    full_text = [paragraph.text for paragraph in doc.paragraphs]
    return '\n'.join(full_text)

def create_document(titulo: str, preguntas: list, buffer: BytesIO, departmento='', examen='', fecha=''):

    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=1.9*cm, leftMargin=1.9*cm, 
                      topMargin=1.9*cm, bottomMargin=1.9*cm)
    document = []
    add_title(document, titulo)
    for preg in preguntas:
        add_paragraph(document, preg['texto'])
        correcta = preg["respuesta_correcta"]
        add_bullets(document, preg['opciones'], correcta)

    doc.build(document)

def add_title(doc: list, title: str, font = 'Helvetica', font_size=36, align= TA_CENTER):
    doc.append(Paragraph(title, ParagraphStyle(name='title', fontFamily=font, fontSize=font_size, alignment=align)))
    doc.append(Spacer(1, 1*cm))

def add_paragraph(doc: list, text: str):
    for line in text.split('\n'):
        doc.append(Paragraph(line))
        doc.append(Spacer(1, 0.5*cm))

def add_bullets(doc: list, lista: list, correcta: str):
    items = []
    for it in lista:
        items.append(ListItem(Paragraph(it)))

    lista_items = ListFlowable(items, bulletType='bullet', start='bulletchar', bulletFontSize=10)

    doc.append(lista_items)
    doc.append(Spacer(1, 0.5*cm))


def add_exam_data(doc: dict, departamento: str, materia: str, profesor: str, fecha: str):
    if departamento != "":
        doc['departamento'] = departamento

    if materia != "":
        doc['materia'] = materia

    if profesor != "":
        doc['profesor'] = profesor
    
    if fecha != "":
        doc['fecha'] = fecha
